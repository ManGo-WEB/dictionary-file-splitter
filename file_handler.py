"""
Модуль для работы с файлами Word
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_dictionary_file(filepath):
    """
    Читает содержимое файла словаря Word
    
    Args:
        filepath (str): Путь к файлу .docx
    Returns:
        Document: Объект документа Word
    """
    try:
        return Document(filepath)
    except Exception as e:
        raise Exception(f"Ошибка при чтении файла: {str(e)}")

def save_section_to_file(letter, paragraphs, output_dir='.'):
    """
    Сохраняет раздел в отдельный файл Word с сохранением форматирования
    
    Args:
        letter (str): Буква раздела
        paragraphs (list): Список параграфов с их форматированием
        output_dir (str): Директория для выходных файлов
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    filename = os.path.join(output_dir, f"{letter}.docx")
    try:
        doc = Document()
        
        # Копируем параграфы с сохранением форматирования
        for text, source_paragraph in paragraphs:
            paragraph = doc.add_paragraph()
            # Копируем текст и его форматирование
            for run in source_paragraph.runs:
                new_run = paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
            # Копируем выравнивание параграфа
            paragraph.alignment = source_paragraph.alignment
            
        doc.save(filename)
        print(f"Создан файл {filename}")
    except Exception as e:
        print(f"Ошибка при сохранении файла {filename}: {str(e)}")